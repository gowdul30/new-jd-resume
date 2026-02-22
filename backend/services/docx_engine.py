"""
DOCX Surgical Swap Engine
Uses python-docx to extract and replace text in Summary and Experience sections
WITHOUT modifying any formatting properties (font, bold, italic, alignment, etc.)
"""

import io
import copy
from typing import Optional
from docx import Document
from docx.oxml.ns import qn

from utils.text_utils import enforce_length_constraint


# Keywords that signal the start of a target section
SUMMARY_KEYWORDS = {"summary", "professional summary", "profile", "objective", "about me", "overview"}
EXPERIENCE_KEYWORDS = {"experience", "work experience", "professional experience", "employment", "career"}
STOP_KEYWORDS = {"education", "skills", "certifications", "projects", "awards", "languages", "references", "volunteer"}


def _para_text_lower(para) -> str:
    return para.text.strip().lower()


def _is_heading(para) -> bool:
    """Check if paragraph is a heading style or bold single-line text."""
    if para.style and para.style.name and "heading" in para.style.name.lower():
        return True
    text = para.text.strip()
    if not text:
        return False
    # Bold single-line paragraph acting as a section header
    if all(run.bold for run in para.runs if run.text.strip()):
        return True
    return False


def extract_docx_sections(file_path_or_bytes) -> dict:
    """
    Reads a DOCX and extracts paragraphs belonging to Summary and Experience sections.

    Returns:
        {
            "summary": [{"para_idx": int, "run_idx": int, "text": str, "section": "summary"}],
            "experience": [...],
            "all_text": str  # full resume text for scoring
        }
    """
    if isinstance(file_path_or_bytes, (str, bytes)):
        doc = Document(file_path_or_bytes)
    else:
        doc = Document(io.BytesIO(file_path_or_bytes))

    sections = {"summary": [], "experience": [], "all_text": ""}
    all_text_parts = []

    current_section = None

    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        all_text_parts.append(text)

        if not text:
            continue

        text_lower = text.lower()

        # Detect section heading
        if _is_heading(para) or (len(text) < 60 and text_lower in SUMMARY_KEYWORDS | EXPERIENCE_KEYWORDS | STOP_KEYWORDS):
            if text_lower in SUMMARY_KEYWORDS:
                current_section = "summary"
            elif text_lower in EXPERIENCE_KEYWORDS:
                current_section = "experience"
            elif text_lower in STOP_KEYWORDS:
                current_section = None
            continue  # Don't capture the heading itself

        if current_section in ("summary", "experience"):
            # Store each meaningful run
            for run_idx, run in enumerate(para.runs):
                if run.text.strip():
                    sections[current_section].append({
                        "para_idx": para_idx,
                        "run_idx": run_idx,
                        "text": run.text,
                    })

    sections["all_text"] = "\n".join(all_text_parts)
    return sections


def inject_docx_rewrites(
    file_path_or_bytes,
    rewrites: dict,  # {"summary": [...new texts per entry...], "experience": [...]}
    tolerance: float = 0.05,
) -> bytes:
    """
    Applies rewritten text to the DOCX, preserving ALL formatting.

    rewrites format:
        {
            "summary": ["new text for run 0", "new text for run 1", ...],
            "experience": [...]
        }

    Returns bytes of the modified DOCX.
    """
    if isinstance(file_path_or_bytes, (str, bytes)):
        doc = Document(file_path_or_bytes)
    else:
        doc = Document(io.BytesIO(file_path_or_bytes))

    # Rebuild sections map to get para/run indices
    sections = extract_docx_sections(file_path_or_bytes)

    for section_name in ("summary", "experience"):
        entries = sections.get(section_name, [])
        new_texts = rewrites.get(section_name, [])

        for i, entry in enumerate(entries):
            if i >= len(new_texts):
                break

            para_idx = entry["para_idx"]
            run_idx = entry["run_idx"]
            original_text = entry["text"]
            new_text = new_texts[i]

            # Enforce ±5% length constraint
            new_text = enforce_length_constraint(original_text, new_text, tolerance)

            # Surgical replacement: only modify .text, nothing else
            para = doc.paragraphs[para_idx]
            run = para.runs[run_idx]
            run.text = new_text  # ONLY this — no font/bold/italic touched

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def get_full_docx_text(file_path_or_bytes) -> str:
    """Returns all text from a DOCX file joined as a single string."""
    if isinstance(file_path_or_bytes, (str, bytes)):
        doc = Document(file_path_or_bytes)
    else:
        doc = Document(io.BytesIO(file_path_or_bytes))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
